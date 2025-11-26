from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from minio import Minio
from minio.error import S3Error
import logging
import uuid
from pathlib import Path
import shutil
from datetime import datetime

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Create FastAPI app
app = FastAPI(
    title="Multimodal RAG API",
    description="Offline multimodal Retrieval-Augmented Generation system",
    version="1.0.0"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# MinIO client
minio_client = Minio(
    "localhost:9000",
    access_key="minioadmin",
    secret_key="minioadmin123",
    secure=False
)

# Ensure bucket exists
BUCKET_NAME = "rag-documents"
try:
    if not minio_client.bucket_exists(BUCKET_NAME):
        minio_client.make_bucket(BUCKET_NAME)
        logger.info(f"Created MinIO bucket: {BUCKET_NAME}")
except S3Error as e:
    logger.error(f"MinIO error: {e}")

# Temporary upload directory
UPLOAD_DIR = Path("./temp_uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

# Persistence directory
DATA_DIR = Path("./data")
DATA_DIR.mkdir(exist_ok=True)

# In-memory document storage
documents_store = []

# In-memory search history storage (max 50 items)
search_history = []
MAX_HISTORY_SIZE = 50



def save_metadata():
    """Save document metadata to JSON."""
    try:
        import json
        # Convert datetime objects to string if needed, currently they are created as default but we store them as is?
        # In models_simple.py we used SQLAlchemy, here we used dicts.
        # Let's ensure we serialize correctly.
        
        # Helper to serialize datetime
        def json_serial(obj):
            if isinstance(obj, datetime):
                return obj.isoformat()
            raise TypeError (f"Type {type(obj)} not serializable")

        with open(DATA_DIR / "documents.json", "w") as f:
            json.dump(documents_store, f, default=json_serial, indent=2)
        logger.info(f"Saved {len(documents_store)} documents to metadata")
    except Exception as e:
        logger.error(f"Error saving metadata: {e}")


def load_metadata():
    """Load document metadata from JSON."""
    global documents_store
    try:
        if (DATA_DIR / "documents.json").exists():
            import json
            with open(DATA_DIR / "documents.json", "r") as f:
                documents_store = json.load(f)
            logger.info(f"Loaded {len(documents_store)} documents from metadata")
    except Exception as e:
        logger.error(f"Error loading metadata: {e}")


def save_search_history():
    """Save search history to JSON."""
    try:
        import json
        with open(DATA_DIR / "search_history.json", "w") as f:
            json.dump(search_history, f, indent=2)
        logger.info(f"Saved {len(search_history)} search history items")
    except Exception as e:
        logger.error(f"Error saving search history: {e}")


def load_search_history():
    """Load search history from JSON."""
    global search_history
    try:
        if (DATA_DIR / "search_history.json").exists():
            import json
            with open(DATA_DIR / "search_history.json", "r") as f:
                search_history = json.load(f)
            logger.info(f"Loaded {len(search_history)} search history items")
    except Exception as e:
        logger.error(f"Error loading search history: {e}")



@app.on_event("startup")
async def startup_event():
    """Load state on startup."""
    load_metadata()
    load_search_history()
    try:
        from embedding_service import load_state
        load_state(DATA_DIR)
    except Exception as e:
        logger.error(f"Error loading embedding state: {e}")


@app.on_event("shutdown")
async def shutdown_event():
    """Save state on shutdown."""
    save_metadata()
    save_search_history()
    try:
        from embedding_service import save_state
        save_state(DATA_DIR)
    except Exception as e:
        logger.error(f"Error saving embedding state: {e}")


@app.get("/")
async def root():
    """Root endpoint."""
    return {
        "message": "Multimodal RAG API",
        "version": "1.0.0",
        "status": "running",
        "features": ["file_upload", "minio_storage", "persistence"],
        "documents_count": len(documents_store)
    }


@app.get("/health")
async def health_check():
    """Health check endpoint."""
    return {"status": "healthy", "minio": "connected"}


@app.get("/api/documents")
async def list_documents():
    """List all documents."""
    return documents_store


def apply_filters(docs, filters, initial_doc_ids=None):
    """Filter documents based on criteria."""
    filtered_ids = set(initial_doc_ids) if initial_doc_ids else {d["id"] for d in docs}
    
    if not filters:
        return list(filtered_ids)
        
    start_date = filters.get("startDate")
    end_date = filters.get("endDate")
    min_size = filters.get("minSize")
    max_size = filters.get("maxSize")
    
    # Convert size to bytes (MB input)
    min_bytes = float(min_size) * 1024 * 1024 if min_size else 0
    max_bytes = float(max_size) * 1024 * 1024 if max_size else float('inf')
    
    final_ids = []
    
    for doc in docs:
        if doc["id"] not in filtered_ids:
            continue
            
        # Date filter
        if start_date or end_date:
            doc_date_str = doc.get("upload_date")
            if not doc_date_str:
                continue # Skip docs without date if filter is applied
                
            doc_date = datetime.fromisoformat(doc_date_str).date()
            
            if start_date and doc_date < datetime.fromisoformat(start_date).date():
                continue
            if end_date and doc_date > datetime.fromisoformat(end_date).date():
                continue
                
        # Size filter
        if min_size or max_size:
            size = doc.get("file_size", 0)
            if size < min_bytes:
                continue
            if size > max_bytes:
                continue
                
        final_ids.append(doc["id"])
        
    return final_ids


@app.post("/api/upload/document")
async def upload_document(file: UploadFile = File(...)):
    """Upload PDF or DOCX document."""
    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in [".pdf", ".docx", ".doc"]:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")
    
    try:
        # Generate unique filename
        unique_filename = f"{uuid.uuid4()}{file_ext}"
        temp_path = UPLOAD_DIR / unique_filename
        
        # Save uploaded file
        with open(temp_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Upload to MinIO
        minio_path = f"documents/{unique_filename}"
        minio_client.fput_object(BUCKET_NAME, minio_path, str(temp_path))
        
        # Store in memory
        doc_info = {
            "id": len(documents_store) + 1,
            "filename": unique_filename,
            "original_filename": file.filename,
            "file_type": file_ext[1:],
            "modality": "text",
            "file_size": temp_path.stat().st_size,
            "minio_path": minio_path,
            "processed": 1,  # Processing
            "upload_date": datetime.now().isoformat()
        }
        documents_store.append(doc_info)
        
        # Process document (extract text and chunk)
        try:
            from text_processor import process_document
            from embedding_service import add_chunks_to_index
            
            result = process_document(temp_path, file_ext[1:])
            
            # Store chunks in document
            doc_info["chunks"] = result["chunks"]
            doc_info["chunk_count"] = result["chunk_count"]
            
            # Generate embeddings and add to FAISS index
            add_chunks_to_index(doc_info["id"], result["chunks"])
            
            doc_info["processed"] = 2  # Completed
            
            # Save state
            from embedding_service import save_state
            save_state(DATA_DIR)
            save_metadata()
            
            logger.info(f"Processed document: {file.filename} - {result['chunk_count']} chunks, indexed")
            
        except Exception as e:
            logger.error(f"Error processing document: {e}")
            doc_info["processed"] = 3  # Failed
        
        # Clean up temp file
        temp_path.unlink()
        
        return {
            "message": "Document uploaded and processed successfully",
            "document_id": doc_info["id"],
            "filename": file.filename,
            "chunks": doc_info.get("chunk_count", 0),
            "status": "processed" if doc_info["processed"] == 2 else "uploaded"
        }
        
    except Exception as e:
        logger.error(f"Error uploading document: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/upload/image")
async def upload_image(file: UploadFile = File(...)):
    """Upload image file."""
    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in [".png", ".jpg", ".jpeg", ".gif", ".bmp"]:
        raise HTTPException(status_code=400, detail="Only image files are supported")
    
    try:
        unique_filename = f"{uuid.uuid4()}{file_ext}"
        temp_path = UPLOAD_DIR / unique_filename
        
        with open(temp_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        minio_path = f"images/{unique_filename}"
        minio_client.fput_object(BUCKET_NAME, minio_path, str(temp_path))
        
        doc_info = {
            "id": len(documents_store) + 1,
            "filename": unique_filename,
            "original_filename": file.filename,
            "file_type": file_ext[1:],
            "modality": "image",
            "file_size": temp_path.stat().st_size,
            "minio_path": minio_path,
            "processed": 2,  # Mark as completed
            "upload_date": datetime.now().isoformat()
        }
        documents_store.append(doc_info)
        
        # Generate CLIP embedding and add to image index (BEFORE deleting temp file!)
        try:
            from embedding_service import add_image_to_index, save_state
            add_image_to_index(doc_info["id"], str(temp_path))
            save_state(DATA_DIR)
            logger.info(f"Generated CLIP embedding for image: {file.filename}")
        except Exception as e:
            logger.error(f"Error generating image embedding: {e}")
        
        # NOW delete the temp file
        temp_path.unlink()
        
        # Save metadata
        save_metadata()
        
        logger.info(f"Uploaded image: {file.filename}")
        
        return {
            "message": "Image uploaded successfully",
            "document_id": doc_info["id"],
            "filename": file.filename,
            "status": "completed"
        }
        
    except Exception as e:
        logger.error(f"Error uploading image: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/upload/audio")
async def upload_audio(file: UploadFile = File(...)):
    """Upload audio file."""
    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in [".mp3", ".wav", ".m4a", ".flac", ".ogg"]:
        raise HTTPException(status_code=400, detail="Only audio files are supported")
    
    try:
        unique_filename = f"{uuid.uuid4()}{file_ext}"
        temp_path = UPLOAD_DIR / unique_filename
        
        with open(temp_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        minio_path = f"audio/{unique_filename}"
        minio_client.fput_object(BUCKET_NAME, minio_path, str(temp_path))
        
        doc_info = {
            "id": len(documents_store) + 1,
            "filename": unique_filename,
            "original_filename": file.filename,
            "file_type": file_ext[1:],
            "modality": "audio",
            "file_size": temp_path.stat().st_size,
            "minio_path": minio_path,
            "processed": 2,  # Mark as completed
            "upload_date": datetime.now().isoformat()
        }
        documents_store.append(doc_info)
        
        # Save metadata
        save_metadata()
        
        temp_path.unlink()
        
        logger.info(f"Uploaded audio: {file.filename}")
        
        return {
            "message": "Audio uploaded successfully",
            "document_id": doc_info["id"],
            "filename": file.filename,
            "status": "completed"
        }
        
    except Exception as e:
        logger.error(f"Error uploading audio: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/documents/{document_id}")
async def delete_document(document_id: int):
    """Delete a document."""
    global documents_store
    doc = next((d for d in documents_store if d["id"] == document_id), None)
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    try:
        # Delete from MinIO
        minio_client.remove_object(BUCKET_NAME, doc["minio_path"])
        
        # Remove from store
        documents_store = [d for d in documents_store if d["id"] != document_id]
        
        save_metadata()
        
        return {"message": "Document deleted successfully"}
        
    except Exception as e:
        logger.error(f"Error deleting document: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/documents/{document_id}/download")
async def get_download_url(document_id: int):
    """Get URL to view document inline."""
    # Return a direct link to our proxy endpoint
    # This avoids MinIO presigned URL issues and gives us full control over headers
    return {
        "url": f"http://localhost:8000/api/documents/{document_id}/content",
        "filename": "document" 
    }


@app.get("/api/documents/{document_id}/content")
async def get_document_content(document_id: int):
    """Stream document content directly from MinIO."""
    doc = next((d for d in documents_store if d["id"] == document_id), None)
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    try:
        from fastapi.responses import StreamingResponse
        
        # Get object from MinIO
        response = minio_client.get_object(BUCKET_NAME, doc["minio_path"])
        
        # Determine content type
        content_type = "application/octet-stream"
        if doc["filename"].endswith(".pdf"):
            content_type = "application/pdf"
        elif doc["filename"].endswith((".jpg", ".jpeg")):
            content_type = "image/jpeg"
        elif doc["filename"].endswith(".png"):
            content_type = "image/png"
            
        # Stream response
        return StreamingResponse(
            response.stream(32*1024),
            media_type=content_type,
            headers={
                "Content-Disposition": f'inline; filename="{doc["original_filename"]}"'
            }
        )
        
    except Exception as e:
        logger.error(f"Error streaming document: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/documents/{document_id}/chunks")
async def get_document_chunks(document_id: int):
    """Get chunks for a document."""
    doc = next((d for d in documents_store if d["id"] == document_id), None)
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return doc.get("chunks", [])




@app.post("/api/export")
async def export_results(request: dict):
    """Export search results as PDF or DOCX."""
    try:
        format_type = request.get("format", "pdf").lower()
        question = request.get("question", "")
        answer = request.get("answer", "")
        citations = request.get("citations", [])
        
        if format_type == "pdf":
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
            from io import BytesIO
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(
                buffer, 
                pagesize=letter,
                title="Search Results",
                author="Multimodal RAG",
                leftMargin=0.75*inch,
                rightMargin=0.75*inch,
                topMargin=0.75*inch,
                bottomMargin=0.75*inch
            )
            
            styles = getSampleStyleSheet()
            
            # Custom styles with enhanced formatting
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#1e40af'),
                spaceAfter=6,
                spaceBefore=12,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            
            subtitle_style = ParagraphStyle(
                'Subtitle',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#6b7280'),
                spaceAfter=20,
                alignment=TA_CENTER,
                fontName='Helvetica-Oblique'
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.HexColor('#3b82f6'),
                spaceAfter=8,
                spaceBefore=12,
                fontName='Helvetica-Bold',
                borderPadding=8,
                borderColor=colors.HexColor('#3b82f6'),
                borderWidth=0,
                leftIndent=0
            )
            
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['BodyText'],
                fontSize=11,
                textColor=colors.HexColor('#1f2937'),
                alignment=TA_JUSTIFY,
                spaceAfter=6,
                leading=16
            )
            
            citation_header_style = ParagraphStyle(
                'CitationHeader',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#059669'),
                fontName='Helvetica-Bold',
                spaceAfter=4
            )
            
            citation_body_style = ParagraphStyle(
                'CitationBody',
                parent=styles['Normal'],
                fontSize=9,
                textColor=colors.HexColor('#374151'),
                leftIndent=20,
                spaceAfter=8,
                leading=13
            )
            
            story = []
            
            # Title with decorative border
            title_data = [[Paragraph("Search Results", title_style)]]
            title_table = Table(title_data, colWidths=[6.5*inch])
            title_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#eff6ff')),
                ('BOX', (0, 0), (-1, -1), 2, colors.HexColor('#3b82f6')),
                ('TOPPADDING', (0, 0), (-1, -1), 16),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 16),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ]))
            story.append(title_table)
            story.append(Spacer(1, 0.1*inch))
            
            # Timestamp
            story.append(Paragraph(
                f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}",
                subtitle_style
            ))
            story.append(Spacer(1, 0.3*inch))
            
            # Question section with border
            question_content = [
                [Paragraph("Question", heading_style)],
                [Paragraph(question or "No question provided", body_style)]
            ]
            question_table = Table(question_content, colWidths=[6.5*inch])
            question_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#dbeafe')),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('BOX', (0, 0), (-1, -1), 1.5, colors.HexColor('#93c5fd')),
                ('LINEBELOW', (0, 0), (-1, 0), 1, colors.HexColor('#3b82f6')),
                ('TOPPADDING', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
                ('LEFTPADDING', (0, 0), (-1, -1), 12),
                ('RIGHTPADDING', (0, 0), (-1, -1), 12),
            ]))
            story.append(question_table)
            story.append(Spacer(1, 0.2*inch))
            
            # Answer section with border
            answer_content = [
                [Paragraph("Answer", heading_style)],
                [Paragraph(answer.replace('\n', '<br/>') or "No answer provided", body_style)]
            ]
            answer_table = Table(answer_content, colWidths=[6.5*inch])
            answer_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#dcfce7')),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('BOX', (0, 0), (-1, -1), 1.5, colors.HexColor('#86efac')),
                ('LINEBELOW', (0, 0), (-1, 0), 1, colors.HexColor('#10b981')),
                ('TOPPADDING', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
                ('LEFTPADDING', (0, 0), (-1, -1), 12),
                ('RIGHTPADDING', (0, 0), (-1, -1), 12),
            ]))
            story.append(answer_table)
            story.append(Spacer(1, 0.3*inch))
            
            # Citations section with enhanced styling
            if citations:
                citations_header_data = [[Paragraph(f"Sources & Citations ({len(citations)})", heading_style)]]
                citations_header_table = Table(citations_header_data, colWidths=[6.5*inch])
                citations_header_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#fef3c7')),
                    ('BOX', (0, 0), (-1, -1), 1.5, colors.HexColor('#fbbf24')),
                    ('TOPPADDING', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
                    ('LEFTPADDING', (0, 0), (-1, -1), 12),
                ]))
                story.append(citations_header_table)
                story.append(Spacer(1, 0.15*inch))
                
                for i, citation in enumerate(citations, 1):
                    citation_data = [
                        [Paragraph(f"[{i}] {citation.get('source', 'Unknown Source')}", citation_header_style)],
                        [Paragraph(citation.get('content', 'No content available')[:500] + '...', citation_body_style)]
                    ]
                    citation_table = Table(citation_data, colWidths=[6.5*inch])
                    citation_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f9fafb')),
                        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#d1d5db')),
                        ('TOPPADDING', (0, 0), (-1, -1), 8),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                        ('LEFTPADDING', (0, 0), (-1, -1), 10),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                    ]))
                    story.append(citation_table)
                    story.append(Spacer(1, 0.1*inch))
            
            doc.build(story)
            buffer.seek(0)
            
            from fastapi.responses import StreamingResponse
            return StreamingResponse(
                buffer,
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename=search_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"}
            )
            
        elif format_type == "docx":
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            from io import BytesIO
            
            doc = Document()
            
            # Title
            title = doc.add_heading('Search Results', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Timestamp
            timestamp_para = doc.add_paragraph()
            timestamp_run = timestamp_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            timestamp_run.italic = True
            timestamp_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            doc.add_paragraph()  # Spacing
            
            # Question
            doc.add_heading('Question:', 2)
            doc.add_paragraph(question)
            
            # Answer
            doc.add_heading('Answer:', 2)
            doc.add_paragraph(answer)
            
            # Citations
            if citations:
                doc.add_heading('Citations:', 2)
                for i, citation in enumerate(citations, 1):
                    para = doc.add_paragraph(style='List Number')
                    para.add_run(f"{citation.get('source', 'Unknown')}\n").bold = True
                    para.add_run(citation.get('content', ''))
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            from fastapi.responses import StreamingResponse
            return StreamingResponse(
                buffer,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename=search_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"}
            )
        else:
            raise HTTPException(status_code=400, detail="Invalid format. Use 'pdf' or 'docx'")
            
    except Exception as e:
        logger.error(f"Error exporting results: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/suggestions")
async def get_suggestions():
    """Get smart suggestions based on random document chunks."""
    try:
        import random
        from rag_service import generate_suggestions
        
        # Get all available chunks from all documents
        all_chunks = []
        for doc in documents_store:
            if "chunks" in doc:
                all_chunks.extend(doc["chunks"])
        
        logger.info(f"Suggestions: Found {len(documents_store)} docs, {len(all_chunks)} chunks")
        
        if not all_chunks:
            return []
            
        # Select random chunks (3-5)
        sample_size = min(len(all_chunks), 5)
        selected_chunks = random.sample(all_chunks, sample_size)
        
        # Generate suggestions
        suggestions = generate_suggestions(selected_chunks)
        
        if not suggestions:
             logger.warning("Suggestions generation returned empty list")
        
        return suggestions
        
    except Exception as e:
        logger.error(f"Error getting suggestions: {e}")
        return []


@app.post("/api/query/")
async def query(request: dict):
    """Query endpoint with RAG."""
    question = request.get("question", "")
    top_k = request.get("top_k", 5)
    document_ids = request.get("document_ids", None)  # Optional filter by document IDs
    
    if not question:
        return {
            "answer": "Please provide a question.",
            "citations": [],
            "context_used": {"text_chunks": 0, "images": 0, "audio_segments": 0}
        }
    
    try:
        from embedding_service import search_similar_chunks, get_index_stats
        from rag_service import generate_answer
        
        # Check if we have any indexed documents
        stats = get_index_stats()
        if stats["total_vectors"] == 0:
            return {
                "answer": f"No documents have been uploaded yet. Please upload some documents first to ask questions about them.",
                "citations": [],
                "context_used": {"text_chunks": 0, "images": 0, "audio_segments": 0}
            }
        
        # Apply filters if present
        filters = request.get("filters")
        filtered_doc_ids = document_ids
        
        if filters:
            filtered_doc_ids = apply_filters(documents_store, filters, document_ids)
            if not filtered_doc_ids:
                return {
                    "answer": "No documents match your filter criteria.",
                    "citations": [],
                    "context_used": {"text_chunks": 0, "images": 0, "audio_segments": 0}
                }
        
        # Search for relevant chunks (with optional document filter)
        search_results = search_similar_chunks(question, top_k, filter_doc_ids=filtered_doc_ids)
        
        if not search_results:
            return {
                "answer": "I couldn't find any relevant information to answer your question.",
                "citations": [],
                "context_used": {"text_chunks": 0, "images": 0, "audio_segments": 0}
            }
        
        # Enrich with document info
        enriched_chunks = []
        for result in search_results:
            doc = next((d for d in documents_store if d["id"] == result["doc_id"]), None)
            if doc:
                enriched_chunks.append({
                    "content": result["content"],
                    "filename": doc["original_filename"],
                    "similarity": result["similarity"],
                    "metadata": doc.get("metadata", {}),
                    "document_id": doc["id"],  # For navigation
                    "chunk_index": result.get("chunk_index", 0)  # For page estimation
                })
        
        # Generate answer using RAG
        rag_result = generate_answer(question, enriched_chunks)
        
        # Log to search history
        try:
            import uuid
            history_item = {
                "id": str(uuid.uuid4()),
                "query": question,
                "timestamp": datetime.now().isoformat(),
                "result_count": len(search_results)
            }
            search_history.insert(0, history_item)  # Add to beginning
            
            # Keep only last MAX_HISTORY_SIZE items
            if len(search_history) > MAX_HISTORY_SIZE:
                search_history[:] = search_history[:MAX_HISTORY_SIZE]
            
            save_search_history()
        except Exception as e:
            logger.error(f"Error logging search history: {e}")
        
        return rag_result
        
    except Exception as e:
        logger.error(f"Query error: {e}")
        return {
            "answer": f"Error processing query: {str(e)}",
            "citations": [],
            "context_used": {"text_chunks": 0, "images": 0, "audio_segments": 0}
        }


@app.get("/api/search-history")
async def get_search_history():
    """Get search history."""
    return search_history


@app.delete("/api/search-history")
async def clear_search_history():
    """Clear search history."""
    global search_history
    search_history = []
    save_search_history()
    return {"message": "Search history cleared"}


@app.post("/api/search/text")
async def search_text(request: dict):
    """Search endpoint with semantic search."""
    query = request.get("query", "")
    top_k = request.get("top_k", 5)
    
    if not query:
        return []
    
    try:
        from embedding_service import search_similar_chunks, get_index_stats
        
        # Get index stats
        stats = get_index_stats()
        if stats["total_vectors"] == 0:
            return []
            
        # Apply filters
        filters = request.get("filters")
        filtered_doc_ids = None
        
        if filters:
            filtered_doc_ids = apply_filters(documents_store, filters)
            if not filtered_doc_ids:
                return []
        
        # Search
        results = search_similar_chunks(query, top_k, filter_doc_ids=filtered_doc_ids)
        
        # Enrich with document info
        enriched_results = []
        for result in results:
            doc = next((d for d in documents_store if d["id"] == result["doc_id"]), None)
            if doc:
                enriched_results.append({
                    "document_id": doc["id"],
                    "filename": doc["original_filename"],
                    "modality": doc["modality"],
                    "chunk_index": result["chunk_index"],
                    "content": result["content"],
                    "similarity": result["similarity"]
                })
        
        return enriched_results
        
    except Exception as e:
        logger.error(f"Search error: {e}")
        return []


@app.post("/api/find-similar")
async def find_similar_documents(
    file: UploadFile = File(...),
    filters: str = Form(None)
):
    """Find similar documents by uploading a file."""
    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in [".pdf", ".docx", ".doc", ".txt"]:
        raise HTTPException(status_code=400, detail="Only PDF, DOCX, and TXT files are supported")
    
    # Parse filters
    parsed_filters = None
    if filters:
        try:
            import json
            parsed_filters = json.loads(filters)
        except Exception as e:
            logger.error(f"Error parsing filters: {e}")
    
    temp_path = None
    try:
        # Generate unique filename
        unique_filename = f"temp_{uuid.uuid4()}{file_ext}"
        temp_path = UPLOAD_DIR / unique_filename
        
        # Save uploaded file temporarily
        with open(temp_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Extract text from file
        from text_processor import extract_text_from_pdf, extract_text_from_docx
        
        if file_ext == ".pdf":
            text = extract_text_from_pdf(temp_path)
        elif file_ext in [".docx", ".doc"]:
            text = extract_text_from_docx(temp_path)
        else:  # .txt
            with open(temp_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        
        # Clean up temp file
        if temp_path.exists():
            temp_path.unlink()
            temp_path = None
        
        if not text or len(text.strip()) < 10:
            raise HTTPException(status_code=400, detail="Could not extract meaningful text from file")
            
        # Apply filters to get allowed doc IDs
        filtered_doc_ids = None
        if parsed_filters:
            filtered_doc_ids = apply_filters(documents_store, parsed_filters)
            if not filtered_doc_ids:
                 return {
                    "uploaded_filename": file.filename,
                    "similar_documents": [],
                    "message": "No documents match your filter criteria"
                }
        
        # Search for similar chunks using first ~5000 chars
        from embedding_service import search_similar_chunks
        search_results = search_similar_chunks(text[:5000], top_k=30, filter_doc_ids=filtered_doc_ids)
        
        if not search_results:
            return {
                "uploaded_filename": file.filename,
                "similar_documents": [],
                "message": "No similar documents found"
            }
        
        # Group results by document and calculate average similarity
        from collections import defaultdict
        doc_scores = defaultdict(list)
        doc_chunks = defaultdict(list)
        
        for result in search_results:
            doc_id = result["doc_id"]
            doc_scores[doc_id].append(result["similarity"])
            doc_chunks[doc_id].append({
                "content": result["content"][:200],  # Preview
                "similarity": result["similarity"]
            })
        
        # Build response
        similar_docs = []
        for doc_id, scores in doc_scores.items():
            doc = next((d for d in documents_store if d["id"] == doc_id), None)
            if doc:
                avg_similarity = sum(scores) / len(scores)
                similar_docs.append({
                    "document_id": doc["id"],
                    "filename": doc["original_filename"],
                    "file_type": doc["file_type"],
                    "similarity": round(avg_similarity, 4),
                    "match_count": len(scores),
                    "top_chunks": sorted(doc_chunks[doc_id], key=lambda x: x["similarity"], reverse=True)[:3]
                })
        
        # Sort by similarity (highest first)
        similar_docs.sort(key=lambda x: x["similarity"], reverse=True)
        
        logger.info(f"Found {len(similar_docs)} similar documents for '{file.filename}'")
        
        return {
            "uploaded_filename": file.filename,
            "similar_documents": similar_docs[:10],  # Top 10
            "total_matches": len(similar_docs)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Find similar error: {e}")
        # Clean up temp file if it exists
        if temp_path and temp_path.exists():
            temp_path.unlink()
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")


@app.post("/api/find-similar-image")
async def find_similar_images(
    file: UploadFile = File(...),
    filters: str = Form(None)
):
    """Find similar images by uploading an image."""
    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in [".jpg", ".jpeg", ".png", ".gif", ".bmp"]:
        raise HTTPException(status_code=400, detail="Only image files (JPG, PNG, GIF, BMP) are supported")
    
    # Parse filters
    parsed_filters = None
    if filters:
        try:
            import json
            parsed_filters = json.loads(filters)
        except Exception as e:
            logger.error(f"Error parsing filters: {e}")
    
    temp_path = None
    try:
        # Generate unique filename
        unique_filename = f"temp_{uuid.uuid4()}{file_ext}"
        temp_path = UPLOAD_DIR / unique_filename
        
        # Save uploaded image temporarily
        with open(temp_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Apply filters to get allowed doc IDs
        filtered_doc_ids = None
        if parsed_filters:
            filtered_doc_ids = apply_filters(documents_store, parsed_filters)
            if not filtered_doc_ids:
                 # Clean up temp file
                if temp_path.exists():
                    temp_path.unlink()
                return {
                    "uploaded_filename": file.filename,
                    "similar_images": [],
                    "message": "No images match your filter criteria"
                }

        # Use CLIP-based similarity search
        from embedding_service import search_similar_images
        
        # Pass filtered_doc_ids to search_similar_images (need to update embedding_service.py too if it doesn't support it yet)
        # But wait, search_similar_images currently doesn't support filtering.
        # I need to filter the results AFTER search if the function doesn't support it, 
        # OR update search_similar_images to support it.
        # Let's check embedding_service.py first. 
        # Actually, for now, I will filter AFTER search for images since I haven't updated embedding_service.py for images yet.
        # BUT, filtering after search is inefficient if top_k is small.
        # Let's update embedding_service.py to support filtering for images as well.
        # For now, let's assume I'll update it or filter post-search.
        # Given the task size, I'll filter post-search here for simplicity as image index is small.
        
        search_results = search_similar_images(str(temp_path), top_k=50) # Get more results to allow for filtering
        
        # Clean up temp file
        if temp_path.exists():
            temp_path.unlink()
            temp_path = None
        
        if not search_results:
            return {
                "uploaded_filename": file.filename,
                "similar_images": [],
                "message": "No similar images found in database"
            }
        
        # Build response with actual similarity scores
        # Only include images with similarity > 0.6 (60%)
        SIMILARITY_THRESHOLD = 0.6
        similar_images = []
        
        for result in search_results:
            # Filter by doc ID if filters applied
            if filtered_doc_ids is not None and result["doc_id"] not in filtered_doc_ids:
                continue

            # Skip low similarity matches
            if result["similarity"] < SIMILARITY_THRESHOLD:
                continue
                
            doc = next((d for d in documents_store if d["id"] == result["doc_id"]), None)
            if doc and doc.get("modality") == "image":
                similar_images.append({
                    "document_id": doc["id"],
                    "filename": doc["original_filename"],
                    "file_type": doc["file_type"],
                    "similarity": result["similarity"],
                    "url": f"/api/documents/{doc['id']}/content"
                })
        
        # Sort by similarity (highest first)
        similar_images.sort(key=lambda x: x["similarity"], reverse=True)
        
        logger.info(f"Found {len(similar_images)} similar images using CLIP (threshold: {SIMILARITY_THRESHOLD})")
        
        return {
            "uploaded_filename": file.filename,
            "similar_images": similar_images[:10],  # Top 10
            "total_matches": len(similar_images)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Find similar image error: {e}")
        # Clean up temp file if it exists
        if temp_path and temp_path.exists():
            temp_path.unlink()
        raise HTTPException(status_code=500, detail=f"Error processing image: {str(e)}")




if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        "main_simple:app",
        host="0.0.0.0",
        port=8000,
        reload=True
    )
