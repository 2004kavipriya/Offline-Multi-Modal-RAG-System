"""
DOCX document processor.
Extracts text from DOCX files with structure preservation.
"""

from docx import Document
from pathlib import Path
from typing import List, Dict, Any
import logging

logger = logging.getLogger(__name__)


class DOCXProcessor:
    """Process DOCX documents and extract text with metadata."""
    
    def __init__(self):
        """Initialize the DOCX processor."""
        pass
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"DOCX file not found: {file_path}")
            
            doc = Document(file_path)
            
            # Extract metadata
            metadata = {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'subject': doc.core_properties.subject or '',
                'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                'modified': str(doc.core_properties.modified) if doc.core_properties.modified else '',
            }
            
            # Extract paragraphs
            paragraphs = []
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    paragraphs.append({
                        'paragraph_number': i + 1,
                        'text': text,
                        'style': para.style.name if para.style else 'Normal'
                    })
            
            # Extract tables
            tables_text = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    tables_text.append({
                        'table_number': table_idx + 1,
                        'data': table_data
                    })
            
            # Combine all text
            full_text = '\n\n'.join([p['text'] for p in paragraphs])
            
            # Add table text
            for table in tables_text:
                table_text = '\n'.join([' | '.join(row) for row in table['data']])
                full_text += f"\n\n[Table {table['table_number']}]\n{table_text}"
            
            logger.info(f"Extracted {len(paragraphs)} paragraphs and {len(tables_text)} tables from {path.name}")
            
            return {
                'text': full_text,
                'paragraphs': paragraphs,
                'tables': tables_text,
                'metadata': metadata,
                'num_paragraphs': len(paragraphs),
                'num_tables': len(tables_text),
                'success': True
            }
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            return {
                'text': '',
                'paragraphs': [],
                'tables': [],
                'metadata': {},
                'num_paragraphs': 0,
                'num_tables': 0,
                'success': False,
                'error': str(e)
            }
    
    def extract_headings(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract headings from the document.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            List of headings with their levels
        """
        try:
            doc = Document(file_path)
            headings = []
            
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    headings.append({
                        'level': level,
                        'text': para.text.strip()
                    })
            
            return headings
            
        except Exception as e:
            logger.error(f"Error extracting headings from {file_path}: {str(e)}")
            return []
